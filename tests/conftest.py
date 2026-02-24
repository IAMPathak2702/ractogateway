"""
Shared pytest fixtures for RactoGateway tests.

Binary test-data files (PDF, DOCX, XLSX, PNG) are created on-the-fly
using the optional library dependencies.  Each fixture is session-scoped
so the files are generated only once per test run.

Run all RAG tests with:
    pip install "ractogateway[rag-all]"
    pytest tests/test_rag_all.py -v
"""
from __future__ import annotations

import contextlib
from pathlib import Path

import pytest

# ---------------------------------------------------------------------------
# Detect optional dependencies
# ---------------------------------------------------------------------------
try:
    import pypdf  # noqa: F401
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import docx  # noqa: F401
    HAS_WORD = True
except ImportError:
    HAS_WORD = False

try:
    import openpyxl  # noqa: F401
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False

try:
    from PIL import Image  # noqa: F401
    HAS_IMAGE = True
except ImportError:
    HAS_IMAGE = False

try:
    import nltk  # noqa: F401
    HAS_NLP = True
except ImportError:
    HAS_NLP = False


# ---------------------------------------------------------------------------
# Reusable skip markers (import these in test files)
# ---------------------------------------------------------------------------
needs_pdf   = pytest.mark.skipif(not HAS_PDF,   reason="pip install ractogateway[rag-pdf]")
needs_word  = pytest.mark.skipif(not HAS_WORD,  reason="pip install ractogateway[rag-word]")
needs_excel = pytest.mark.skipif(not HAS_EXCEL, reason="pip install ractogateway[rag-excel]")
needs_image = pytest.mark.skipif(not HAS_IMAGE, reason="pip install ractogateway[rag-image]")
needs_nlp   = pytest.mark.skipif(not HAS_NLP,   reason="pip install ractogateway[rag-nlp]")


# ---------------------------------------------------------------------------
# Static data directory
# ---------------------------------------------------------------------------
DATA_DIR = Path(__file__).parent / "data"


# ---------------------------------------------------------------------------
# Binary fixture: PDF
# ---------------------------------------------------------------------------
@pytest.fixture(scope="session")
def pdf_path(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """Return path to a minimal two-page PDF test file.

    Requires pypdf (``pip install ractogateway[rag-pdf]``).
    If not installed the path still points to a valid location but the
    PdfReader tests will be skipped by their ``needs_pdf`` marker.
    """
    path: Path = tmp_path_factory.mktemp("bindata") / "sample.pdf"
    if not HAS_PDF:
        return path

    # Build a minimal valid PDF without needing reportlab
    pdf_content = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj

2 0 obj
<< /Type /Pages /Kids [3 0 R 5 0 R] /Count 2 >>
endobj

3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
   /Contents 4 0 R /Resources << /Font << /F1 7 0 R >> >> >>
endobj

4 0 obj
<< /Length 120 >>
stream
BT
/F1 14 Tf
72 720 Td
(RactoGateway PDF Test - Page 1) Tj
0 -30 Td
(Testing PdfReader with pypdf library.) Tj
0 -30 Td
(Python is great for AI development.) Tj
ET
endstream
endobj

5 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
   /Contents 6 0 R /Resources << /Font << /F1 7 0 R >> >> >>
endobj

6 0 obj
<< /Length 100 >>
stream
BT
/F1 14 Tf
72 720 Td
(RactoGateway PDF Test - Page 2) Tj
0 -30 Td
(Machine learning transforms every industry.) Tj
ET
endstream
endobj

7 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj

xref
0 8
0000000000 65535 f
0000000009 00000 n
0000000068 00000 n
0000000125 00000 n
0000000274 00000 n
0000000446 00000 n
0000000597 00000 n
0000000749 00000 n

trailer
<< /Size 8 /Root 1 0 R >>
startxref
820
%%EOF
"""
    path.write_bytes(pdf_content)
    return path


# ---------------------------------------------------------------------------
# Binary fixture: DOCX
# ---------------------------------------------------------------------------
@pytest.fixture(scope="session")
def docx_path(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """Return path to a minimal .docx Word document.

    Requires python-docx (``pip install ractogateway[rag-word]``).
    """
    path: Path = tmp_path_factory.mktemp("bindata") / "sample.docx"
    if not HAS_WORD:
        return path

    from docx import Document as DocxDocument  # type: ignore[import]

    doc = DocxDocument()
    doc.add_heading("RactoGateway Word Test Document", level=0)
    doc.add_paragraph(
        "This document is used to test the WordReader class. "
        "The reader extracts text from all paragraphs in the .docx file."
    )
    doc.add_heading("Section 1: Overview", level=1)
    doc.add_paragraph(
        "RactoGateway provides a unified interface to OpenAI, Google Gemini, "
        "and Anthropic Claude models."
    )
    doc.add_heading("Section 2: RAG Components", level=1)
    doc.add_paragraph("The RAG pipeline consists of readers, chunkers, processors, "
                       "embedders, and vector stores.")
    for item in ["TextReader", "PdfReader", "WordReader", "SpreadsheetReader"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Section 3: Conclusion", level=1)
    doc.add_paragraph(
        "With RactoGateway you can build production-grade AI applications "
        "without worrying about provider-specific APIs."
    )
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Binary fixture: XLSX
# ---------------------------------------------------------------------------
@pytest.fixture(scope="session")
def xlsx_path(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """Return path to a minimal .xlsx Excel workbook with two sheets.

    Requires openpyxl (``pip install ractogateway[rag-excel]``).
    """
    path: Path = tmp_path_factory.mktemp("bindata") / "sample.xlsx"
    if not HAS_EXCEL:
        return path

    import openpyxl  # type: ignore[import]

    wb = openpyxl.Workbook()

    # Sheet 1 — product catalog
    ws1 = wb.active
    assert ws1 is not None
    ws1.title = "Products"
    ws1.append(["Name", "Category", "Price", "Rating", "In Stock"])
    ws1.append(["Laptop Pro", "Electronics", 1299.99, 4.8, True])
    ws1.append(["Wireless Mouse", "Accessories", 29.99, 4.5, True])
    ws1.append(["USB Hub", "Accessories", 49.99, 4.3, True])
    ws1.append(["Mechanical Keyboard", "Peripherals", 149.99, 4.7, False])

    # Sheet 2 — employee roster
    ws2 = wb.create_sheet(title="Employees")
    ws2.append(["Name", "Department", "Role", "Salary"])
    ws2.append(["Alice Chen", "Engineering", "Senior Engineer", 120000])
    ws2.append(["Bob Smith", "Product", "Product Manager", 110000])
    ws2.append(["Carol Jones", "Data Science", "ML Engineer", 125000])

    wb.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Binary fixture: PNG
# ---------------------------------------------------------------------------
@pytest.fixture(scope="session")
def png_path(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """Return path to a small PNG image with coloured rectangles.

    Requires Pillow (``pip install ractogateway[rag-image]``).
    """
    path: Path = tmp_path_factory.mktemp("bindata") / "sample.png"
    if not HAS_IMAGE:
        return path

    from PIL import Image as PILImage  # type: ignore[import]
    from PIL import ImageDraw, ImageFont

    img = PILImage.new("RGB", (256, 128), color=(30, 30, 30))
    draw = ImageDraw.Draw(img)
    # Coloured rectangles
    draw.rectangle([10, 10, 80, 60],  fill=(100, 149, 237))   # cornflower blue
    draw.rectangle([90, 10, 160, 60], fill=(255, 165, 0))     # orange
    draw.rectangle([170, 10, 240, 60], fill=(60, 179, 113))   # medium sea green
    # Try to add text (font may not be available everywhere)
    with contextlib.suppress(Exception):
        draw.text((10, 75), "RactoGateway Image Test", fill=(255, 255, 255),
                  font=ImageFont.load_default())
    img.save(str(path), format="PNG")
    return path


# ---------------------------------------------------------------------------
# Binary fixture: JPEG
# ---------------------------------------------------------------------------
@pytest.fixture(scope="session")
def jpg_path(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """Return path to a small JPEG image.

    Requires Pillow (``pip install ractogateway[rag-image]``).
    """
    path: Path = tmp_path_factory.mktemp("bindata") / "sample.jpg"
    if not HAS_IMAGE:
        return path

    from PIL import Image as PILImage  # type: ignore[import]

    img = PILImage.new("RGB", (128, 128), color=(220, 50, 47))  # red gradient
    img.save(str(path), format="JPEG", quality=85)
    return path
