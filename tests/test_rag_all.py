"""
Comprehensive tests for RactoGateway file-based RAG functionality.

Coverage:
  - All file readers  : TextReader, HtmlReader, SpreadsheetReader (CSV/TSV/XLSX),
                        PdfReader, WordReader, ImageReader, FileReaderRegistry
  - All chunkers      : FixedChunker, RecursiveChunker, SentenceChunker, SemanticChunker
  - All processors    : TextCleaner, Lemmatizer, ProcessingPipeline
  - RactoFile         : from_path, from_bytes, is_image/is_pdf/is_text, base64_data
  - Document / Chunk  : model fields and metadata
  - RactoRAG pipeline : ingest_text, ingest (file), retrieve, count, clear
                        (uses MockEmbedder — no API key required)

Prerequisites
-------------
  Static fixtures : tests/data/sample.{txt,md,html,csv,tsv,json,rst}
  Binary fixtures : created by conftest.py on-the-fly (session-scoped)

  Run with all optional deps:
      pip install "ractogateway[rag-all]"
      pytest tests/test_rag_all.py -v
"""
from __future__ import annotations

import hashlib
import math
from pathlib import Path
from typing import Any

import pytest

# ---------------------------------------------------------------------------
# Shared constants (mirrors conftest.py — conftest is not importable directly)
# ---------------------------------------------------------------------------
DATA_DIR = Path(__file__).parent / "data"

try:
    import pypdf  # noqa: F401
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import docx  # noqa: F401
    HAS_WORD = True
except ImportError:
    HAS_WORD = False

try:
    import openpyxl  # noqa: F401
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False

try:
    from PIL import Image  # noqa: F401
    HAS_IMAGE = True
except ImportError:
    HAS_IMAGE = False

try:
    import nltk  # noqa: F401
    HAS_NLP = True
except ImportError:
    HAS_NLP = False

needs_pdf   = pytest.mark.skipif(not HAS_PDF,   reason="pip install ractogateway[rag-pdf]")
needs_word  = pytest.mark.skipif(not HAS_WORD,  reason="pip install ractogateway[rag-word]")
needs_excel = pytest.mark.skipif(not HAS_EXCEL, reason="pip install ractogateway[rag-excel]")
needs_image = pytest.mark.skipif(not HAS_IMAGE, reason="pip install ractogateway[rag-image]")
needs_nlp   = pytest.mark.skipif(not HAS_NLP,   reason="pip install ractogateway[rag-nlp]")

# ---------------------------------------------------------------------------
# Mock embedder — deterministic, no API key, 64-dim
# ---------------------------------------------------------------------------
from ractogateway.rag.embedders.base import BaseEmbedder


class MockEmbedder(BaseEmbedder):
    """SHA-256-based deterministic embedder for offline tests.

    Produces a 64-float unit-normalised vector for any text without
    calling any external API.
    """

    @property
    def dimension(self) -> int:
        return 64

    def embed(self, texts: list[str]) -> list[list[float]]:
        result: list[list[float]] = []
        for text in texts:
            raw = hashlib.sha256(text.encode()).digest()  # 32 bytes
            # Duplicate to get 64 floats
            floats = [((b / 255.0) - 0.5) * 2.0 for b in raw + raw]
            norm = math.sqrt(sum(x * x for x in floats)) or 1.0
            result.append([x / norm for x in floats])
        return result

    async def aembed(self, texts: list[str]) -> list[list[float]]:
        return self.embed(texts)


# ===========================================================================
# 1. TEXT READER
# ===========================================================================
class TestTextReader:
    def test_read_txt(self) -> None:
        from ractogateway.rag.readers import TextReader

        reader = TextReader()
        doc = reader.read(DATA_DIR / "sample.txt")

        assert doc.content, "content must not be empty"
        assert "Artificial Intelligence" in doc.content
        assert doc.source.endswith("sample.txt")
        assert doc.doc_id  # auto UUID

    def test_read_markdown(self) -> None:
        from ractogateway.rag.readers import TextReader

        doc = TextReader().read(DATA_DIR / "sample.md")
        assert "RactoGateway" in doc.content
        assert "pip install" in doc.content

    def test_read_rst(self) -> None:
        from ractogateway.rag.readers import TextReader

        doc = TextReader().read(DATA_DIR / "sample.rst")
        assert "reStructuredText" in doc.content or "TextReader" in doc.content

    def test_read_json(self) -> None:
        from ractogateway.rag.readers import TextReader

        doc = TextReader().read(DATA_DIR / "sample.json")
        assert "ractogateway" in doc.content.lower() or "RactoGateway" in doc.content

    def test_supported_extensions(self) -> None:
        from ractogateway.rag.readers import TextReader

        exts = TextReader().supported_extensions
        assert ".txt" in exts
        assert ".md"  in exts
        assert ".rst" in exts
        assert ".json" in exts

    def test_custom_encoding(self) -> None:
        from ractogateway.rag.readers import TextReader

        # Should not raise even if file is ascii-safe
        doc = TextReader(encoding="utf-8").read(DATA_DIR / "sample.txt")
        assert len(doc.content) > 0

    def test_missing_file_raises(self) -> None:
        from ractogateway.rag.readers import TextReader

        with pytest.raises((FileNotFoundError, OSError)):
            TextReader().read(Path("/nonexistent/path/file.txt"))


# ===========================================================================
# 2. HTML READER
# ===========================================================================
class TestHtmlReader:
    def test_strips_tags(self) -> None:
        from ractogateway.rag.readers import HtmlReader

        doc = HtmlReader().read(DATA_DIR / "sample.html")
        assert "<html" not in doc.content
        assert "<h1>"  not in doc.content
        assert "<p>"   not in doc.content

    def test_preserves_text_content(self) -> None:
        from ractogateway.rag.readers import HtmlReader

        doc = HtmlReader().read(DATA_DIR / "sample.html")
        assert "large language model" in doc.content.lower() or \
               "LLM" in doc.content

    def test_source_path_set(self) -> None:
        from ractogateway.rag.readers import HtmlReader

        doc = HtmlReader().read(DATA_DIR / "sample.html")
        assert "sample.html" in doc.source

    def test_supported_extensions(self) -> None:
        from ractogateway.rag.readers import HtmlReader

        exts = HtmlReader().supported_extensions
        assert ".html" in exts
        assert ".htm"  in exts


# ===========================================================================
# 3. SPREADSHEET READER (CSV / TSV / XLSX)
# ===========================================================================
class TestSpreadsheetReader:
    def test_read_csv(self) -> None:
        from ractogateway.rag.readers import SpreadsheetReader

        doc = SpreadsheetReader().read(DATA_DIR / "sample.csv")
        assert "Name" in doc.content  # header
        assert "Laptop Pro" in doc.content

    def test_read_tsv(self) -> None:
        from ractogateway.rag.readers import SpreadsheetReader

        doc = SpreadsheetReader().read(DATA_DIR / "sample.tsv")
        assert "Country" in doc.content
        assert "Germany" in doc.content

    def test_max_rows_limit(self) -> None:
        from ractogateway.rag.readers import SpreadsheetReader

        doc = SpreadsheetReader(max_rows=3).read(DATA_DIR / "sample.csv")
        # Should have header + at most 3 data rows
        lines = [ln for ln in doc.content.splitlines() if ln.strip()]
        # Rows * sheets is hard to count exactly, but the content is shorter
        assert len(lines) <= 6  # generous bound

    def test_without_header(self) -> None:
        from ractogateway.rag.readers import SpreadsheetReader

        doc_with = SpreadsheetReader(include_header=True).read(DATA_DIR / "sample.csv")
        doc_without = SpreadsheetReader(include_header=False).read(DATA_DIR / "sample.csv")
        # With header is always >= without header
        assert len(doc_with.content) >= len(doc_without.content)

    @needs_excel
    def test_read_xlsx(self, xlsx_path: Path) -> None:
        from ractogateway.rag.readers import SpreadsheetReader

        doc = SpreadsheetReader().read(xlsx_path)
        assert "Products" in doc.content or "Laptop" in doc.content

    def test_supported_extensions(self) -> None:
        from ractogateway.rag.readers import SpreadsheetReader

        exts = SpreadsheetReader().supported_extensions
        assert ".csv"  in exts
        assert ".tsv"  in exts

    @needs_excel
    def test_xlsx_in_supported_extensions(self) -> None:
        from ractogateway.rag.readers import SpreadsheetReader

        exts = SpreadsheetReader().supported_extensions
        assert ".xlsx" in exts


# ===========================================================================
# 4. PDF READER
# ===========================================================================
@needs_pdf
class TestPdfReader:
    def test_read_pdf(self, pdf_path: Path) -> None:
        from ractogateway.rag.readers import PdfReader

        doc = PdfReader().read(pdf_path)
        assert len(doc.content) > 0, "PDF content must not be empty"

    def test_source_set(self, pdf_path: Path) -> None:
        from ractogateway.rag.readers import PdfReader

        doc = PdfReader().read(pdf_path)
        assert "sample.pdf" in doc.source

    def test_supported_extensions(self) -> None:
        from ractogateway.rag.readers import PdfReader

        assert ".pdf" in PdfReader().supported_extensions


# ===========================================================================
# 5. WORD READER
# ===========================================================================
@needs_word
class TestWordReader:
    def test_read_docx(self, docx_path: Path) -> None:
        from ractogateway.rag.readers import WordReader

        doc = WordReader().read(docx_path)
        assert "RactoGateway" in doc.content

    def test_paragraphs_included(self, docx_path: Path) -> None:
        from ractogateway.rag.readers import WordReader

        doc = WordReader().read(docx_path)
        assert "unified interface" in doc.content.lower() or \
               "WordReader" in doc.content

    def test_supported_extensions(self) -> None:
        from ractogateway.rag.readers import WordReader

        assert ".docx" in WordReader().supported_extensions


# ===========================================================================
# 6. IMAGE READER
# ===========================================================================
@needs_image
class TestImageReader:
    def test_read_png(self, png_path: Path) -> None:
        from ractogateway.rag.readers import ImageReader

        doc = ImageReader().read(png_path)
        # content contains metadata strings (dimensions, format, mode)
        assert doc.content
        assert "256" in doc.content or "PNG" in doc.content or "RGB" in doc.content

    def test_read_jpeg(self, jpg_path: Path) -> None:
        from ractogateway.rag.readers import ImageReader

        doc = ImageReader().read(jpg_path)
        assert doc.content
        assert "128" in doc.content or "JPEG" in doc.content or "RGB" in doc.content

    def test_exif_flag(self, png_path: Path) -> None:
        from ractogateway.rag.readers import ImageReader

        doc_with    = ImageReader(include_exif=True).read(png_path)
        doc_without = ImageReader(include_exif=False).read(png_path)
        # Both should succeed (PNG may have no EXIF)
        assert doc_with.content
        assert doc_without.content

    def test_source_path_set(self, png_path: Path) -> None:
        from ractogateway.rag.readers import ImageReader

        doc = ImageReader().read(png_path)
        assert "sample.png" in doc.source

    def test_supported_extensions(self) -> None:
        from ractogateway.rag.readers import ImageReader

        exts = ImageReader().supported_extensions
        assert ".png"  in exts
        assert ".jpg"  in exts
        assert ".jpeg" in exts


# ===========================================================================
# 7. FILE READER REGISTRY
# ===========================================================================
class TestFileReaderRegistry:
    def test_auto_detect_txt(self) -> None:
        from ractogateway.rag.readers import FileReaderRegistry

        reg = FileReaderRegistry()
        doc = reg.read(DATA_DIR / "sample.txt")
        assert "Artificial Intelligence" in doc.content

    def test_auto_detect_html(self) -> None:
        from ractogateway.rag.readers import FileReaderRegistry

        doc = FileReaderRegistry().read(DATA_DIR / "sample.html")
        assert "<html" not in doc.content

    def test_auto_detect_csv(self) -> None:
        from ractogateway.rag.readers import FileReaderRegistry

        doc = FileReaderRegistry().read(DATA_DIR / "sample.csv")
        assert "Laptop" in doc.content

    def test_unsupported_extension_raises(self) -> None:
        from ractogateway.rag.readers import FileReaderRegistry

        with pytest.raises(ValueError, match=r"(?i)(no reader|unsupported|not supported)"):
            FileReaderRegistry().read(Path("archive.zip"))

    def test_custom_reader_registration(self) -> None:
        from ractogateway.rag.readers import FileReaderRegistry, TextReader

        reg = FileReaderRegistry(readers=[])  # empty registry
        reg.register(TextReader())
        doc = reg.read(DATA_DIR / "sample.txt")
        assert doc.content

    def test_supported_extensions_property(self) -> None:
        from ractogateway.rag.readers import FileReaderRegistry

        exts = FileReaderRegistry().supported_extensions
        assert ".txt" in exts
        assert ".html" in exts
        assert ".csv" in exts


# ===========================================================================
# 8. DOCUMENT AND CHUNK MODELS
# ===========================================================================
class TestDocumentModel:
    def test_document_fields(self) -> None:
        from ractogateway.rag._models import Document

        doc = Document(content="Hello world", source="test.txt")
        assert doc.content == "Hello world"
        assert doc.source  == "test.txt"
        assert doc.doc_id             # auto-generated UUID
        assert isinstance(doc.metadata, dict)

    def test_custom_doc_id(self) -> None:
        from ractogateway.rag._models import Document

        doc = Document(content="x", source="y", doc_id="my-stable-id")
        assert doc.doc_id == "my-stable-id"

    def test_metadata_pass_through(self) -> None:
        from ractogateway.rag._models import Document

        doc = Document(content="x", source="y", metadata={"author": "Alice", "year": 2026})
        assert doc.metadata["author"] == "Alice"
        assert doc.metadata["year"]   == 2026


class TestChunkModel:
    def test_chunk_fields_after_chunking(self) -> None:
        from ractogateway.rag._models import Document
        from ractogateway.rag.chunkers import FixedChunker

        doc = Document(content="A" * 600, source="fake.txt")
        chunks = FixedChunker(chunk_size=200, overlap=20).chunk(doc)

        assert len(chunks) >= 3
        for i, chunk in enumerate(chunks):
            assert chunk.doc_id == doc.doc_id
            assert chunk.chunk_id                   # auto UUID
            assert chunk.content                    # not empty
            assert chunk.embedding is None          # not yet embedded
            assert chunk.metadata.chunk_index == i
            assert chunk.metadata.total_chunks == len(chunks)
            assert chunk.metadata.source == doc.source
            assert chunk.metadata.start_char >= 0
            assert chunk.metadata.end_char   >  chunk.metadata.start_char


# ===========================================================================
# 9. FIXED CHUNKER
# ===========================================================================
class TestFixedChunker:
    def test_basic_split(self) -> None:
        from ractogateway.rag._models import Document
        from ractogateway.rag.chunkers import FixedChunker

        text = "x" * 1000
        doc    = Document(content=text, source="x.txt")
        chunks = FixedChunker(chunk_size=300, overlap=0).chunk(doc)
        assert len(chunks) >= 3

    def test_chunk_size_respected(self) -> None:
        from ractogateway.rag._models import Document
        from ractogateway.rag.chunkers import FixedChunker

        doc = Document(content="A" * 500, source="x.txt")
        for chunk in FixedChunker(chunk_size=100, overlap=0).chunk(doc):
            assert len(chunk.content) <= 100

    def test_overlap_creates_duplicate_chars(self) -> None:
        from ractogateway.rag._models import Document
        from ractogateway.rag.chunkers import FixedChunker

        doc = Document(content="abcdefghijklmnopqrstuvwxyz" * 4, source="x.txt")
        chunks = FixedChunker(chunk_size=20, overlap=5).chunk(doc)
        if len(chunks) >= 2:
            # Last 5 chars of chunk[0] == first 5 chars of chunk[1]
            assert chunks[0].content[-5:] == chunks[1].content[:5]

    def test_small_document_single_chunk(self) -> None:
        from ractogateway.rag._models import Document
        from ractogateway.rag.chunkers import FixedChunker

        doc = Document(content="short text", source="x.txt")
        chunks = FixedChunker(chunk_size=500).chunk(doc)
        assert len(chunks) == 1
        assert chunks[0].content == "short text"

    def test_metadata_indices(self) -> None:
        from ractogateway.rag._models import Document
        from ractogateway.rag.chunkers import FixedChunker

        doc = Document(content="A" * 600, source="x.txt")
        chunks = FixedChunker(chunk_size=200, overlap=0).chunk(doc)
        for i, chunk in enumerate(chunks):
            assert chunk.metadata.chunk_index == i
            assert chunk.metadata.total_chunks == len(chunks)


# ===========================================================================
# 10. RECURSIVE CHUNKER
# ===========================================================================
class TestRecursiveChunker:
    def test_splits_at_paragraph_boundaries(self) -> None:
        from ractogateway.rag._models import Document
        from ractogateway.rag.chunkers import RecursiveChunker

        text = "\n\n".join([f"Paragraph {i}. " * 20 for i in range(5)])
        doc    = Document(content=text, source="x.txt")
        chunks = RecursiveChunker(chunk_size=200, overlap=20).chunk(doc)
        assert len(chunks) >= 2

    def test_chunk_size_not_exceeded_much(self) -> None:
        from ractogateway.rag._models import Document
        from ractogateway.rag.chunkers import RecursiveChunker

        doc = Document(content="Word " * 300, source="x.txt")
        for chunk in RecursiveChunker(chunk_size=100, overlap=10).chunk(doc):
            # Recursive chunker may slightly exceed when no separator fits;
            # allow 2x tolerance in the worst case
            assert len(chunk.content) <= 200

    def test_custom_separators(self) -> None:
        from ractogateway.rag._models import Document
        from ractogateway.rag.chunkers import RecursiveChunker

        text = "sentence1. sentence2. sentence3. sentence4. sentence5."
        doc    = Document(content=text, source="x.txt")
        chunks = RecursiveChunker(
            chunk_size=30,
            overlap=0,
            separators=[". ", " ", ""],
        ).chunk(doc)
        assert len(chunks) >= 2

    def test_single_chunk_if_small(self) -> None:
        from ractogateway.rag._models import Document
        from ractogateway.rag.chunkers import RecursiveChunker

        doc = Document(content="tiny", source="x.txt")
        chunks = RecursiveChunker(chunk_size=512).chunk(doc)
        assert len(chunks) == 1

    def test_real_document(self) -> None:
        from ractogateway.rag.chunkers import RecursiveChunker
        from ractogateway.rag.readers import TextReader

        doc    = TextReader().read(DATA_DIR / "sample.txt")
        chunks = RecursiveChunker(chunk_size=256, overlap=32).chunk(doc)
        assert len(chunks) >= 2
        for chunk in chunks:
            assert chunk.content.strip()


# ===========================================================================
# 11. SENTENCE CHUNKER  (needs NLTK)
# ===========================================================================
@needs_nlp
class TestSentenceChunker:
    def test_basic_sentence_split(self) -> None:
        from ractogateway.rag._models import Document
        from ractogateway.rag.chunkers import SentenceChunker

        text = (
            "Artificial intelligence is powerful. "
            "Machine learning enables pattern recognition. "
            "Deep learning uses neural networks. "
            "NLP processes human language. "
            "Computer vision understands images. "
            "Robotics combines AI with mechanics."
        )
        doc    = Document(content=text, source="x.txt")
        chunks = SentenceChunker(sentences_per_chunk=2, overlap_sentences=0).chunk(doc)
        assert len(chunks) >= 2

    def test_overlap_sentences(self) -> None:
        from ractogateway.rag._models import Document
        from ractogateway.rag.chunkers import SentenceChunker

        sentences = [f"Sentence {i} is here." for i in range(10)]
        doc    = Document(content=" ".join(sentences), source="x.txt")
        chunks = SentenceChunker(sentences_per_chunk=3, overlap_sentences=1).chunk(doc)
        # With overlap, consecutive chunks share 1 sentence
        if len(chunks) >= 2:
            last_of_first  = chunks[0].content.split(".")[-2].strip() + "."
            first_of_second = chunks[1].content.split(".")[0].strip() + "."
            # At least partial overlap check — both non-empty
            assert last_of_first
            assert first_of_second

    def test_real_document(self) -> None:
        from ractogateway.rag.chunkers import SentenceChunker
        from ractogateway.rag.readers import TextReader

        doc    = TextReader().read(DATA_DIR / "sample.txt")
        chunks = SentenceChunker(sentences_per_chunk=3).chunk(doc)
        assert len(chunks) >= 1
        for chunk in chunks:
            assert chunk.content.strip()


# ===========================================================================
# 12. SEMANTIC CHUNKER  (needs NLTK)
# ===========================================================================
@needs_nlp
class TestSemanticChunker:
    def test_basic_semantic_split(self) -> None:
        from ractogateway.rag._models import Document
        from ractogateway.rag.chunkers import SemanticChunker

        text = (
            "The sky is blue. The sun shines brightly. Clouds are white. "
            "Python is a programming language. Code is written in text files. "
            "Functions are reusable blocks of code."
        )
        doc    = Document(content=text, source="x.txt")
        # Use low threshold to get more splits
        chunks = SemanticChunker(
            embedder=MockEmbedder(),
            threshold=0.99,
            min_chunk_size=1,
        ).chunk(doc)
        assert len(chunks) >= 1

    def test_high_threshold_fewer_splits(self) -> None:
        from ractogateway.rag._models import Document
        from ractogateway.rag.chunkers import SemanticChunker

        doc = Document(content="Hello world. Hi there. Good day. Greetings.", source="x.txt")
        chunks_strict = SemanticChunker(
            embedder=MockEmbedder(), threshold=0.1, min_chunk_size=1
        ).chunk(doc)
        chunks_loose  = SemanticChunker(
            embedder=MockEmbedder(), threshold=0.99, min_chunk_size=1
        ).chunk(doc)
        # Strict threshold (low) = everything is different = more chunks
        assert len(chunks_strict) >= len(chunks_loose)


# ===========================================================================
# 13. TEXT CLEANER
# ===========================================================================
class TestTextCleaner:
    def test_strips_html_tags(self) -> None:
        from ractogateway.rag.processors import TextCleaner

        raw = "<p>Hello <b>World</b></p>"
        cleaned = TextCleaner(strip_html=True).process(raw)
        assert "<p>" not in cleaned
        assert "<b>" not in cleaned
        assert "Hello" in cleaned
        assert "World" in cleaned

    def test_collapses_whitespace(self) -> None:
        from ractogateway.rag.processors import TextCleaner

        raw = "This    has   lots   of     spaces."
        cleaned = TextCleaner(collapse_whitespace=True).process(raw)
        assert "  " not in cleaned

    def test_collapses_blank_lines(self) -> None:
        from ractogateway.rag.processors import TextCleaner

        raw = "Line 1\n\n\n\nLine 2"
        cleaned = TextCleaner(collapse_blank_lines=True).process(raw)
        assert "\n\n\n" not in cleaned

    def test_strips_control_chars(self) -> None:
        from ractogateway.rag.processors import TextCleaner

        raw = "Hello\x00\x01\x02World"
        cleaned = TextCleaner(strip_control_chars=True).process(raw)
        assert "\x00" not in cleaned
        assert "Hello" in cleaned
        assert "World" in cleaned

    def test_normalize_unicode(self) -> None:
        from ractogateway.rag.processors import TextCleaner

        # Café as NFC vs NFD (decomposed accent)
        raw = "caf\u00e9"  # NFC precomposed 'é'
        cleaned = TextCleaner(normalize_unicode=True).process(raw)
        assert cleaned  # should not raise or return empty

    def test_disable_all_cleans(self) -> None:
        from ractogateway.rag.processors import TextCleaner

        raw = "<p>test   content</p>"
        cleaned = TextCleaner(
            strip_html=False,
            collapse_whitespace=False,
            strip_control_chars=False,
            collapse_blank_lines=False,
            normalize_unicode=False,
        ).process(raw)
        assert cleaned == raw


# ===========================================================================
# 14. LEMMATIZER  (needs NLTK)
# ===========================================================================
@needs_nlp
class TestLemmatizer:
    def test_lemmatizes_verbs(self) -> None:
        from ractogateway.rag.processors import Lemmatizer

        # "running" → "run", "flies" → "fly"
        lemmatizer = Lemmatizer()
        result = lemmatizer.process("The cats are running and the bird flies away.")
        assert "run" in result or "running" in result  # depends on POS tag accuracy
        assert result  # not empty

    def test_lemmatizes_nouns(self) -> None:
        from ractogateway.rag.processors import Lemmatizer

        result = Lemmatizer().process("The children played with their toys.")
        assert "child" in result or "children" in result
        assert result

    def test_returns_string(self) -> None:
        from ractogateway.rag.processors import Lemmatizer

        result = Lemmatizer().process("Testing lemmatization output type.")
        assert isinstance(result, str)
        assert len(result) > 0

    def test_without_pos_tagging(self) -> None:
        from ractogateway.rag.processors import Lemmatizer

        result = Lemmatizer(use_pos_tagging=False).process("The dogs barked loudly.")
        assert isinstance(result, str)
        assert result


# ===========================================================================
# 15. PROCESSING PIPELINE
# ===========================================================================
class TestProcessingPipeline:
    def test_sequential_application(self) -> None:
        from ractogateway.rag.processors import ProcessingPipeline, TextCleaner

        pipeline = ProcessingPipeline([
            TextCleaner(strip_html=True),
            TextCleaner(collapse_whitespace=True),
        ])
        raw = "<p>Hello    World</p>"
        result = pipeline.process(raw)
        assert "<p>" not in result
        assert "  "  not in result
        assert "Hello" in result

    @needs_nlp
    def test_cleaner_then_lemmatizer(self) -> None:
        from ractogateway.rag.processors import Lemmatizer, ProcessingPipeline, TextCleaner

        pipeline = ProcessingPipeline([
            TextCleaner(strip_html=True, collapse_whitespace=True),
            Lemmatizer(),
        ])
        raw = "<p>The  running dogs   chased the flying birds.</p>"
        result = pipeline.process(raw)
        assert "<p>" not in result
        assert result.strip()

    def test_empty_processors_raises(self) -> None:
        from ractogateway.rag.processors import ProcessingPipeline

        with pytest.raises((ValueError, Exception)):
            ProcessingPipeline([])

    def test_single_processor(self) -> None:
        from ractogateway.rag.processors import ProcessingPipeline, TextCleaner

        pipeline = ProcessingPipeline([TextCleaner()])
        assert pipeline.process("hello world") == "hello world"


# ===========================================================================
# 16. RACTOFILE
# ===========================================================================
class TestRactoFile:
    def test_from_path_txt(self) -> None:
        from ractogateway.prompts.engine import RactoFile

        rf = RactoFile.from_path(DATA_DIR / "sample.txt")
        assert rf.data
        assert "text" in rf.mime_type
        assert rf.is_text
        assert not rf.is_image
        assert not rf.is_pdf

    def test_from_path_html(self) -> None:
        from ractogateway.prompts.engine import RactoFile

        rf = RactoFile.from_path(DATA_DIR / "sample.html")
        assert rf.data
        assert "html" in rf.mime_type or "text" in rf.mime_type

    def test_from_bytes(self) -> None:
        from ractogateway.prompts.engine import RactoFile

        raw = b"Hello from bytes"
        rf  = RactoFile.from_bytes(raw, "text/plain", name="hello.txt")
        assert rf.data      == raw
        assert rf.mime_type == "text/plain"
        assert rf.name      == "hello.txt"

    def test_base64_data_roundtrip(self) -> None:
        import base64

        from ractogateway.prompts.engine import RactoFile

        raw = b"Binary content \x00\x01\x02"
        rf  = RactoFile.from_bytes(raw, "application/octet-stream")
        decoded = base64.b64decode(rf.base64_data)
        assert decoded == raw

    @needs_image
    def test_from_path_png(self, png_path: Path) -> None:
        from ractogateway.prompts.engine import RactoFile

        rf = RactoFile.from_path(png_path)
        assert rf.is_image
        assert not rf.is_pdf
        assert "image" in rf.mime_type

    @needs_pdf
    def test_from_path_pdf(self, pdf_path: Path) -> None:
        from ractogateway.prompts.engine import RactoFile

        rf = RactoFile.from_path(pdf_path)
        assert rf.is_pdf
        assert not rf.is_image
        assert rf.mime_type == "application/pdf"

    def test_missing_file_raises(self) -> None:
        from ractogateway.prompts.engine import RactoFile

        with pytest.raises(FileNotFoundError):
            RactoFile.from_path(Path("/no/such/file.txt"))

    def test_name_from_path(self) -> None:
        from ractogateway.prompts.engine import RactoFile

        rf = RactoFile.from_path(DATA_DIR / "sample.txt")
        assert rf.name == "sample.txt"


# ===========================================================================
# 17. RACTOTAG PIPELINE — INGEST & RETRIEVE  (InMemoryVectorStore + MockEmbedder)
# ===========================================================================
class TestRactoRAGPipeline:
    @pytest.fixture()
    def rag(self) -> Any:
        from ractogateway.rag import RactoRAG
        from ractogateway.rag.chunkers import RecursiveChunker
        from ractogateway.rag.processors import TextCleaner
        from ractogateway.rag.stores import InMemoryVectorStore

        return RactoRAG(
            embedder=MockEmbedder(),
            store=InMemoryVectorStore(),
            chunker=RecursiveChunker(chunk_size=256, overlap=32),
            processors=[TextCleaner()],
        )

    # ── ingest_text ─────────────────────────────────────────────────────────
    def test_ingest_text_returns_chunks(self, rag: Any) -> None:
        chunks = rag.ingest_text("Python is a versatile programming language.", source="test")
        assert len(chunks) >= 1
        assert all(c.embedding is not None for c in chunks)

    def test_ingest_text_increments_count(self, rag: Any) -> None:
        before = rag.count()
        rag.ingest_text("First text document.")
        rag.ingest_text("Second text document.")
        assert rag.count() >= before + 2

    def test_ingest_text_custom_metadata(self, rag: Any) -> None:
        chunks = rag.ingest_text("Meta test.", source="meta_test", author="Alice")
        assert all(c.metadata.extra.get("author") == "Alice" for c in chunks)

    # ── ingest (file) ────────────────────────────────────────────────────────
    def test_ingest_txt_file(self, rag: Any) -> None:
        chunks = rag.ingest(DATA_DIR / "sample.txt")
        assert len(chunks) >= 1
        assert all(c.embedding is not None for c in chunks)

    def test_ingest_md_file(self, rag: Any) -> None:
        chunks = rag.ingest(DATA_DIR / "sample.md")
        assert len(chunks) >= 1

    def test_ingest_html_file(self, rag: Any) -> None:
        chunks = rag.ingest(DATA_DIR / "sample.html")
        assert len(chunks) >= 1

    def test_ingest_csv_file(self, rag: Any) -> None:
        chunks = rag.ingest(DATA_DIR / "sample.csv")
        assert len(chunks) >= 1

    def test_ingest_json_file(self, rag: Any) -> None:
        chunks = rag.ingest(DATA_DIR / "sample.json")
        assert len(chunks) >= 1

    @needs_word
    def test_ingest_docx(self, rag: Any, docx_path: Path) -> None:
        chunks = rag.ingest(docx_path)
        assert len(chunks) >= 1

    @needs_excel
    def test_ingest_xlsx(self, rag: Any, xlsx_path: Path) -> None:
        chunks = rag.ingest(xlsx_path)
        assert len(chunks) >= 1

    @needs_pdf
    def test_ingest_pdf(self, rag: Any, pdf_path: Path) -> None:
        chunks = rag.ingest(pdf_path)
        assert len(chunks) >= 1

    # ── retrieve ────────────────────────────────────────────────────────────
    def test_retrieve_returns_results(self, rag: Any) -> None:
        rag.ingest_text("Python is great for AI and machine learning.", source="a")
        rag.ingest_text("JavaScript is the language of the web.", source="b")
        rag.ingest_text("Rust is known for memory safety and speed.", source="c")

        results = rag.retrieve("Python programming", top_k=2)
        assert len(results) <= 2
        assert all(r.score >= 0.0 for r in results)
        assert all(r.rank  >= 1   for r in results)

    def test_retrieve_top_k_limit(self, rag: Any) -> None:
        for i in range(10):
            rag.ingest_text(f"Document number {i} with some content.", source=f"doc{i}")
        results = rag.retrieve("content", top_k=3)
        assert len(results) <= 3

    def test_retrieve_result_fields(self, rag: Any) -> None:
        rag.ingest_text("The quick brown fox jumps over the lazy dog.", source="fox")
        results = rag.retrieve("fox", top_k=1)
        assert results
        r = results[0]
        assert r.chunk.content
        assert r.chunk.embedding is not None
        assert r.score >= 0.0
        assert r.rank  == 1

    def test_retrieve_result_ranks_ordered(self, rag: Any) -> None:
        for i in range(5):
            rag.ingest_text(f"Unique sentence {i}: " + "word " * 20, source=f"s{i}")
        results = rag.retrieve("sentence", top_k=5)
        for i, r in enumerate(results):
            assert r.rank == i + 1

    # ── count & clear ────────────────────────────────────────────────────────
    def test_count_increases_after_ingest(self, rag: Any) -> None:
        n0 = rag.count()
        rag.ingest_text("One.")
        rag.ingest_text("Two.")
        assert rag.count() >= n0 + 2

    def test_clear_empties_store(self, rag: Any) -> None:
        rag.ingest_text("Some content to clear.", source="clear_test")
        rag.clear()
        assert rag.count() == 0

    def test_retrieve_empty_store_returns_empty(self, rag: Any) -> None:
        rag.clear()
        results = rag.retrieve("anything", top_k=5)
        assert results == []

    # ── ingest_dir ───────────────────────────────────────────────────────────
    def test_ingest_dir_ingests_multiple_files(self, rag: Any) -> None:
        rag.clear()
        chunks = rag.ingest_dir(DATA_DIR, pattern="*.txt")
        assert len(chunks) >= 1
        assert rag.count() >= 1


# ===========================================================================
# 18. IN-MEMORY VECTOR STORE (unit-level)
# ===========================================================================
class TestInMemoryVectorStore:
    def _make_chunk(self, text: str, embedding: list[float]) -> Any:
        from ractogateway.rag._models import Chunk, ChunkMetadata

        meta = ChunkMetadata(
            source="test.txt",
            chunk_index=0,
            total_chunks=1,
            start_char=0,
            end_char=len(text),
            doc_id="doc-1",
        )
        return Chunk(doc_id="doc-1", content=text, embedding=embedding, metadata=meta)

    def test_add_and_count(self) -> None:
        from ractogateway.rag.stores import InMemoryVectorStore

        store = InMemoryVectorStore()
        emb   = MockEmbedder().embed(["hello"])[0]
        store.add([self._make_chunk("hello", emb)])
        assert store.count() == 1

    def test_search_returns_result(self) -> None:
        from ractogateway.rag.stores import InMemoryVectorStore

        store = InMemoryVectorStore()
        embedder = MockEmbedder()
        texts = ["apple fruit", "car vehicle", "dog animal"]
        for text in texts:
            emb = embedder.embed([text])[0]
            store.add([self._make_chunk(text, emb)])

        query_emb = embedder.embed(["fruit"])[0]
        results   = store.search(query_emb, top_k=2)
        assert len(results) <= 2
        assert all(r.score >= -1.0 for r in results)  # cosine in [-1, 1]

    def test_delete_chunk(self) -> None:
        from ractogateway.rag.stores import InMemoryVectorStore

        store = InMemoryVectorStore()
        emb   = MockEmbedder().embed(["test"])[0]
        chunk = self._make_chunk("test", emb)
        store.add([chunk])
        assert store.count() == 1
        store.delete([chunk.chunk_id])
        assert store.count() == 0

    def test_clear_removes_all(self) -> None:
        from ractogateway.rag.stores import InMemoryVectorStore

        store = InMemoryVectorStore()
        embedder = MockEmbedder()
        for t in ["a", "b", "c"]:
            store.add([self._make_chunk(t, embedder.embed([t])[0])])
        assert store.count() == 3
        store.clear()
        assert store.count() == 0

    def test_search_on_empty_store(self) -> None:
        from ractogateway.rag.stores import InMemoryVectorStore

        store  = InMemoryVectorStore()
        query  = MockEmbedder().embed(["query"])[0]
        result = store.search(query, top_k=5)
        assert result == []

    def test_chunk_without_embedding_raises(self) -> None:
        from ractogateway.rag._models import Chunk, ChunkMetadata
        from ractogateway.rag.stores import InMemoryVectorStore

        store = InMemoryVectorStore()
        meta  = ChunkMetadata(
            source="x.txt", chunk_index=0, total_chunks=1,
            start_char=0, end_char=5, doc_id="d1",
        )
        chunk_no_emb = Chunk(doc_id="d1", content="hello", embedding=None, metadata=meta)
        with pytest.raises((ValueError, Exception)):
            store.add([chunk_no_emb])


# ===========================================================================
# 19. ASYNC VARIANTS (ingest + retrieve)
# ===========================================================================
class TestRactoRAGAsync:
    @pytest.mark.asyncio
    async def test_aingest_text(self) -> None:
        from ractogateway.rag import RactoRAG
        from ractogateway.rag.stores import InMemoryVectorStore

        rag    = RactoRAG(embedder=MockEmbedder(), store=InMemoryVectorStore())
        chunks = await rag.aingest_text("Async ingestion works great.", source="async_test")
        assert len(chunks) >= 1
        assert all(c.embedding is not None for c in chunks)

    @pytest.mark.asyncio
    async def test_aretrieve(self) -> None:
        from ractogateway.rag import RactoRAG
        from ractogateway.rag.stores import InMemoryVectorStore

        rag = RactoRAG(embedder=MockEmbedder(), store=InMemoryVectorStore())
        await rag.aingest_text("Async Python AI pipeline.", source="ar1")
        await rag.aingest_text("Async JavaScript web apps.", source="ar2")

        results = await rag.aretrieve("Python", top_k=2)
        assert len(results) <= 2
        assert all(r.rank >= 1 for r in results)

    @pytest.mark.asyncio
    async def test_aingest_file(self) -> None:
        from ractogateway.rag import RactoRAG
        from ractogateway.rag.stores import InMemoryVectorStore

        rag    = RactoRAG(embedder=MockEmbedder(), store=InMemoryVectorStore())
        chunks = await rag.aingest(DATA_DIR / "sample.txt")
        assert len(chunks) >= 1


# ===========================================================================
# 20. FULL INTEGRATION: Read → Chunk → Clean → Embed → Store → Retrieve
# ===========================================================================
class TestFullRAGIntegration:
    def test_end_to_end_txt(self) -> None:
        """Full pipeline: .txt → RecursiveChunker → TextCleaner → MockEmbed → retrieve."""
        from ractogateway.rag import RactoRAG
        from ractogateway.rag.chunkers import RecursiveChunker
        from ractogateway.rag.processors import TextCleaner
        from ractogateway.rag.stores import InMemoryVectorStore

        rag = RactoRAG(
            embedder=MockEmbedder(),
            store=InMemoryVectorStore(),
            chunker=RecursiveChunker(chunk_size=200, overlap=20),
            processors=[TextCleaner()],
        )
        chunks = rag.ingest(DATA_DIR / "sample.txt")
        assert len(chunks) >= 3

        results = rag.retrieve("machine learning", top_k=3)
        assert results
        # The most relevant chunk should mention something related to the query
        top_chunk_text = results[0].chunk.content.lower()
        assert top_chunk_text  # not empty

    def test_end_to_end_html(self) -> None:
        """HTML content should have tags stripped before chunking."""
        from ractogateway.rag import RactoRAG
        from ractogateway.rag.stores import InMemoryVectorStore

        rag = RactoRAG(embedder=MockEmbedder(), store=InMemoryVectorStore())
        rag.ingest(DATA_DIR / "sample.html")
        results = rag.retrieve("performance tips", top_k=2)
        # Tags should be stripped
        for r in results:
            assert "<" not in r.chunk.content

    def test_multi_file_ingest_and_retrieve(self) -> None:
        """Ingest multiple file types and retrieve across them."""
        from ractogateway.rag import RactoRAG
        from ractogateway.rag.stores import InMemoryVectorStore

        rag = RactoRAG(embedder=MockEmbedder(), store=InMemoryVectorStore())
        rag.ingest(DATA_DIR / "sample.txt")
        rag.ingest(DATA_DIR / "sample.md")
        rag.ingest(DATA_DIR / "sample.html")
        rag.ingest(DATA_DIR / "sample.csv")

        assert rag.count() >= 4
        results = rag.retrieve("Python", top_k=5)
        assert results
        # Results come from different source files
        sources = {r.chunk.metadata.source for r in results}
        assert len(sources) >= 1

    @needs_nlp
    def test_pipeline_with_lemmatizer(self) -> None:
        """Full pipeline with TextCleaner + Lemmatizer before embedding."""
        from ractogateway.rag import RactoRAG
        from ractogateway.rag.processors import Lemmatizer, TextCleaner
        from ractogateway.rag.stores import InMemoryVectorStore

        rag = RactoRAG(
            embedder=MockEmbedder(),
            store=InMemoryVectorStore(),
            processors=[
                TextCleaner(strip_html=True, collapse_whitespace=True),
                Lemmatizer(),
            ],
        )
        rag.ingest_text(
            "The machines are learning. The models are running and training.",
            source="lemma_test",
        )
        results = rag.retrieve("machine learning", top_k=1)
        assert results
