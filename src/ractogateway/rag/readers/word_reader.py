"""Word document reader — uses ``python-docx`` (lazy import).

Install with:  pip install ractogateway[rag-word]
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any


def _require_docx() -> Any:
    try:
        import docx
    except ImportError as exc:
        raise ImportError(
            "Reading Word (.docx) files requires the 'python-docx' package. "
            "Install it with:  pip install ractogateway[rag-word]"
        ) from exc
    return docx


from ractogateway.rag._models.document import Document
from ractogateway.rag.readers.base import BaseReader


class WordReader(BaseReader):
    """Extract text from Microsoft Word (.docx) files using ``python-docx``.

    Accepts a file path (``str`` / ``Path``), raw ``bytes``, or any binary
    file-like object with a ``.read()`` method.
    """

    @property
    def supported_extensions(self) -> frozenset[str]:
        return frozenset({".docx"})

    def _read_path(self, path: Path) -> Document:
        docx = _require_docx()
        doc = docx.Document(str(path))
        paragraphs, tables_text = self._extract_content(doc)
        full_text = "\n\n".join(paragraphs + tables_text)
        core_props = doc.core_properties
        return Document(
            content=full_text,
            source=str(path.resolve()),
            metadata={
                "extension": ".docx",
                "filename": path.name,
                "size_bytes": path.stat().st_size,
                "author": core_props.author or "",
                "title": core_props.title or "",
                "paragraph_count": len(paragraphs),
            },
        )

    def _read_bytes(self, data: bytes, *, source_label: str = "<bytes>") -> Document:
        docx = _require_docx()
        doc = docx.Document(io.BytesIO(data))
        paragraphs, tables_text = self._extract_content(doc)
        full_text = "\n\n".join(paragraphs + tables_text)
        core_props = doc.core_properties
        return Document(
            content=full_text,
            source=source_label,
            metadata={
                "size_bytes": len(data),
                "author": core_props.author or "",
                "title": core_props.title or "",
                "paragraph_count": len(paragraphs),
            },
        )

    def _extract_content(self, doc: Any) -> tuple[list[str], list[str]]:
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_text: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    tables_text.append(row_text)
        return paragraphs, tables_text
